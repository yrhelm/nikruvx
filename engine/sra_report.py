"""
HIPAA Security Risk Analysis (SRA) generator.

Produces a markdown evidence pack auto-populated from the graph (and
optionally a Word docx if `python-docx` is installed). Aligned to the HHS
"Guidance on Risk Analysis" structure - the core sections covering organizational
inventory, threats, vulnerabilities, existing controls, gaps, remediation.

Usage:
    from engine.sra_report import generate_markdown, generate_docx
    md = generate_markdown(stack_summary="My PHI app", scope="Production")
"""
from __future__ import annotations

import datetime as _dt
from io import BytesIO

from .graph import run_read
from .healthcare import list_phi_packages_in_graph
from .hipaa import hipaa_coverage, hipaa_gaps_for_cve
from .posture import gaps_for_cve as _gaps_for_cve


def _today() -> str:
    return _dt.date.today().isoformat()


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------
def _section_assets() -> str:
    pkgs = list_phi_packages_in_graph()
    if not pkgs:
        return ("_No PHI-handling packages currently identified in the graph. "
                "Run `python -c 'from engine.healthcare import tag_phi_packages; "
                "tag_phi_packages()'` after ingesting your SBOM._")
    rows = "\n".join(
        f"| {p['ecosystem']} | {p['name']} | {p['cve_count']} |"
        for p in pkgs[:50]
    )
    return (
        "| Ecosystem | Package | CVEs affecting |\n"
        "|---|---|---|\n" + rows
    )


def _section_threats_and_vulns() -> str:
    rows = run_read("""
        MATCH (c:CVE)-[:AFFECTS]->(p:Package:HandlesPHI)
        OPTIONAL MATCH (c)-[:MAPS_TO]->(l:OSILayer)
        OPTIONAL MATCH (c)-[:CLASSIFIED_AS]->(w:CWE)
        RETURN c.id AS id, c.severity AS sev, c.cvss_score AS cvss,
               collect(DISTINCT l.number) AS layers,
               collect(DISTINCT w.id)[0..3] AS cwes,
               c.description AS desc
        ORDER BY coalesce(c.cvss_score,0) DESC LIMIT 40
    """)
    if not rows:
        return "_No PHI-affecting CVEs in the graph yet._"
    out = ["| CVE | Severity | CVSS | OSI Layers | CWE | Summary |",
           "|---|---|---|---|---|---|"]
    for r in rows:
        out.append(
            f"| {r['id']} | {r['sev'] or '?'} | {r['cvss'] or '—'} "
            f"| {','.join('L'+str(x) for x in r['layers'] if x) or '—'} "
            f"| {','.join(r['cwes']) or '—'} "
            f"| {(r['desc'] or '').replace('|',' ').strip()[:140]}… |"
        )
    return "\n".join(out)


def _section_controls() -> str:
    cov = hipaa_coverage()
    if not cov["regulatory_capabilities"]:
        return "_No policy controls loaded. Use the Posture tab to upload your IAM/SG/CA/Intune/firewall policies._"
    out = [
        "| Capability | Coverage | Controls in graph | HIPAA / GDPR / FDA Citations |",
        "|---|---|---|---|",
    ]
    for c in cov["regulatory_capabilities"]:
        cites = "; ".join(f"{x['framework']} §{x['section']}" for x in c["citations"])
        out.append(
            f"| **{c['capability']}** | {c['coverage_pct']}% | {c['controls_total']} | {cites} |"
        )
    return "\n".join(out)


def _section_gaps_and_remediation() -> str:
    """For each PHI-affecting CVE, run the gap analyzer and dump the hottest gaps."""
    cves = run_read("""
        MATCH (c:CVE)-[:AFFECTS]->(p:Package:HandlesPHI)
        RETURN c.id AS id ORDER BY coalesce(c.cvss_score,0) DESC LIMIT 8
    """)
    if not cves:
        return "_No PHI-affecting CVEs to analyze._"
    sections = []
    for r in cves:
        cve_id = r["id"]
        gaps = hipaa_gaps_for_cve(cve_id)
        if not gaps.get("gaps"):
            continue
        s = [f"### {cve_id}", ""]
        s.append("| Step | Capability | Severity | Required Controls | Citations |")
        s.append("|---|---|---|---|---|")
        for g in gaps["gaps"][:5]:
            cites = "; ".join(f"{c['framework']} §{c['section']}" for c in g.get("citations", []))
            classes = ", ".join(g.get("recommended_classes", [])[:3])
            s.append(f"| L{g['step_layer']} | {g['capability']} | {g['severity']} | {classes} | {cites} |")
        sections.append("\n".join(s))
    return "\n\n".join(sections) if sections else "_All analyzed PHI CVEs are covered by your loaded policies._"


def _section_existing_blocks() -> str:
    cves = run_read("""
        MATCH (c:CVE)-[:AFFECTS]->(p:Package:HandlesPHI)
        RETURN c.id AS id ORDER BY coalesce(c.cvss_score,0) DESC LIMIT 5
    """)
    if not cves:
        return "_No data._"
    s = ["| CVE | Capability blocked | Control | Source |", "|---|---|---|---|"]
    seen = set()
    for r in cves:
        g = _gaps_for_cve(r["id"])
        for b in g.get("blocks", [])[:3]:
            for c in (b.get("controls") or [])[:1]:
                key = (r["id"], b["capability"], c.get("id"))
                if key in seen: continue
                seen.add(key)
                s.append(f"| {r['id']} | {b['capability']} | {c.get('title','?')} | {c.get('source','?')} |")
    return "\n".join(s) if len(s) > 2 else "_No active blocks yet - upload your policies._"


# ---------------------------------------------------------------------------
# Public API: markdown + docx
# ---------------------------------------------------------------------------
def generate_markdown(stack_summary: str = "", scope: str = "Production environment",
                       organization: str = "Your Organization") -> str:
    cov = hipaa_coverage()
    return f"""# HIPAA Security Risk Analysis
**Organization:** {organization}
**Scope:** {scope}
**Date:** {_today()}
**Generated by:** Cyber Nexus

---

## Executive Summary

This analysis was auto-generated from the Cyber Nexus graph - a real-time
inventory of your software components, the vulnerabilities affecting them,
the security policies you have deployed, and the gaps between them. It
satisfies the four core sections of an HHS-aligned **Security Risk Analysis**
under HIPAA Security Rule §164.308(a)(1)(ii)(A): identification of assets
handling ePHI, threats and vulnerabilities to those assets, current security
measures, and a remediation plan.

**Stack Description:**
{stack_summary or "_Provide a stack description when triggering the report._"}

**At a glance:**
- PHI-handling components identified: **{cov['phi_package_count']}**
- CVEs affecting PHI components: **{cov['phi_cve_count']}**
- Regulatory capabilities tracked: **{len(cov['regulatory_capabilities'])}**

---

## 1. Assets Handling Electronic PHI (164.308(a)(1)(ii)(A))

The following software components in your environment process, store, or
transmit ePHI as detected by Cyber Nexus PHI-package signatures
(HL7v2 / FHIR / DICOM / X12 / clinical NLP / EHR client SDKs / biosignals).

{_section_assets()}

---

## 2. Threats and Vulnerabilities (164.308(a)(1)(ii)(A) - Risk Analysis)

CVEs affecting your PHI-handling components, ranked by CVSS. OSI layer column
shows where each weakness expresses itself; CWE column shows the weakness
class.

{_section_threats_and_vulns()}

---

## 3. Current Security Measures (164.308(a)(1)(ii)(B) / 164.312)

Capability coverage from the policies you have uploaded into Cyber Nexus
(IAM, Security Groups, Conditional Access, Intune, WAF, firewall, etc.),
mapped to the HIPAA Security Rule sections each control class addresses.

{_section_controls()}

### Active blocks (where existing controls intercept attack capabilities)

{_section_existing_blocks()}

---

## 4. Risk Determination & Remediation Plan (164.308(a)(1)(ii)(B))

For each high-priority PHI-affecting CVE, the gap analyzer enumerated the
attack capabilities not covered by your existing policies. Each row maps to
the specific HIPAA / GDPR / FDA citation it implicates.

{_section_gaps_and_remediation()}

---

## 5. Documentation Standard (164.316(b)(2))

This document is regenerated automatically each time it is requested - the
underlying graph is the source of truth. Maintain version snapshots in your
GRC repository and re-run after any material change to the SBOM, IAM/CA/SG
policies, or upon disclosure of a new CVE affecting a PHI component.

**Cyber Nexus generates this report from a graph database. The graph is the
evidence; this document is its narrative.**

---

_End of Risk Analysis_
"""


def generate_docx(stack_summary: str = "", scope: str = "Production environment",
                   organization: str = "Your Organization") -> bytes:
    """Optional docx variant. Falls back to a stub message if python-docx absent."""
    try:
        from docx import Document
    except ImportError:
        md = generate_markdown(stack_summary, scope, organization)
        return ("# (python-docx not installed — fall back to markdown)\n\n" + md).encode("utf-8")

    md = generate_markdown(stack_summary, scope, organization)
    doc = Document()

    # Lightweight markdown -> docx pass: headings, paragraphs, tables.
    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        s = line.rstrip()
        if s.startswith("# "):
            doc.add_heading(s[2:], 0)
        elif s.startswith("## "):
            doc.add_heading(s[3:], 1)
        elif s.startswith("### "):
            doc.add_heading(s[4:], 2)
        elif s.startswith("|") and i + 1 < len(lines) and lines[i+1].startswith("|"):
            # Table block
            table_rows = []
            while i < len(lines) and lines[i].startswith("|"):
                row = [c.strip() for c in lines[i].strip("|").split("|")]
                table_rows.append(row); i += 1
            i -= 1
            if len(table_rows) >= 2:
                table = doc.add_table(rows=1, cols=len(table_rows[0]))
                hdr = table.rows[0].cells
                for c, val in zip(hdr, table_rows[0]):
                    c.text = val.replace("**","")
                # Skip the separator row (---|---)
                for row in table_rows[2:]:
                    cells = table.add_row().cells
                    for c, val in zip(cells, row):
                        c.text = val.replace("**","").replace("`","")
        elif s.startswith("---"):
            pass
        elif s.startswith("**") and s.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(s.strip("*"))
            run.bold = True
        elif s:
            doc.add_paragraph(s)
        i += 1

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
