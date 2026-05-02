"""
Model Card + FDA SaMD Readiness Generator
==========================================
Produces a markdown (and optional docx) Model Card for any AI/ML model in
your environment, populated from:
   - Cyber Nexus graph (CVEs in deps, AI threats linked, posture gaps)
   - Clinical AI test runner findings
   - User-provided narrative bits (intended use, training data, etc.)

Sections follow the FDA Good Machine Learning Practice (GMLP) 10 Guiding
Principles + the FDA Pre-Submission package layout.
"""

from __future__ import annotations

import datetime as _dt
from io import BytesIO

from .clinical_runner import list_findings, summary
from .graph import run_read
from .hipaa import owasp_llm_citations


def _today() -> str:
    return _dt.date.today().isoformat()


def _ai_threats_summary() -> str:
    rows = run_read("""
        MATCH (a:AIThreat)
        RETURN a.id AS id, a.name AS name, a.framework AS framework
        ORDER BY a.framework, a.id
    """)
    if not rows:
        return "_No AI threat catalog loaded._"
    lines = ["| ID | Framework | Name | HIPAA / FDA citations |", "|---|---|---|---|"]
    for r in rows[:25]:
        cites = "; ".join(f"{c['framework']} §{c['section']}" for c in owasp_llm_citations(r["id"]))
        lines.append(f"| {r['id']} | {r['framework']} | {r['name']} | {cites or '—'} |")
    return "\n".join(lines)


def _findings_table(model: str | None) -> tuple[str, dict]:
    s = summary(model)
    findings = list_findings(model, limit=200)
    if not findings:
        return (
            "_No clinical-AI test results yet. Run `python -m engine.clinical_runner --model "
            f"{model or '<your-model>'}`._",
            s,
        )
    by_cat: dict[str, list] = {}
    for f in findings:
        by_cat.setdefault(f["category"], []).append(f)
    parts = []
    for cat, items in by_cat.items():
        passed = sum(1 for f in items if f["passed"])
        failed = len(items) - passed
        parts.append(f"\n### {cat}  ({passed} pass, {failed} fail)\n")
        parts.append("| Test | Passed | Severity | Capability | Reason | Citation |")
        parts.append("|---|---|---|---|---|---|")
        for f in items[:8]:
            parts.append(
                f"| {f['test_id']} | {'✓' if f['passed'] else '✗'} "
                f"| {f.get('severity', '—')} | {f.get('capability', '—')} "
                f"| {(f.get('reason') or '').replace('|', ' ')[:80]} "
                f"| {(f.get('citation') or '').replace('|', ' ')} |"
            )
    return "\n".join(parts), s


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------
def generate_markdown(
    model_name: str = "(unspecified)",
    intended_use: str = "",
    training_data: str = "",
    evaluation_data: str = "",
    limitations: str = "",
    monitoring_plan: str = "",
    deployment_context: str = "",
    organization: str = "Your Organization",
) -> str:
    findings_md, summ = _findings_table(model_name)
    pass_total = summ["total_pass"]
    fail_total = summ["total_fail"]
    overall = (
        "READY"
        if fail_total == 0 and pass_total > 0
        else "NEEDS_WORK"
        if fail_total
        else "UNTESTED"
    )
    return f"""# Model Card — {model_name}
**Organization:** {organization}
**Generated:** {_today()}
**Source:** Cyber Nexus (graph + clinical-AI test runner)
**Overall readiness:** **{overall}**  (passed {pass_total} / failed {fail_total})

---

## 1. Intended Use  (FDA GMLP Principle 1)

{intended_use or "_Describe the clinical task, target population, deployment setting, and clinical decision support level (advisory vs. autonomous)._"}

## 2. Deployment Context  (GMLP Principle 6)

{deployment_context or "_Describe where in the workflow the model is invoked, who reviews its output, what guardrails are in place._"}

## 3. Training Data  (GMLP Principles 2, 3)

{training_data or "_Describe data sources, time range, demographic distribution, sites, BAA status, de-identification standard (Safe Harbor / Expert Determination), label provenance._"}

## 4. Evaluation Data and Performance  (GMLP Principles 3, 4, 8)

{evaluation_data or "_Describe held-out evaluation set, demographic strata, primary and secondary metrics, confidence intervals, performance vs. existing standard of care._"}

## 5. Adversarial / Safety Test Results  (GMLP Principle 5)

The following adversarial tests were run from the Cyber Nexus clinical-AI
test corpus. Categories cover drug confusion, dose injection, ICD
manipulation, indirect prompt injection through clinical artifacts,
de-identification reversal, demographic bias, hallucinated guidelines, and
output safety.

{findings_md}

## 6. AI Threat Surface  (GMLP Principle 5 + OWASP LLM Top 10)

The following adversarial-ML risks are tracked in the graph for this model
and its dependencies; each is mapped to the relevant HIPAA / GDPR / FDA
citation.

{_ai_threats_summary()}

## 7. Known Limitations and Out-of-Scope Use  (GMLP Principle 9)

{limitations or "_Describe populations, settings, languages, image modalities, edge cases for which the model is NOT validated. Forbid use cases out of scope._"}

## 8. Monitoring Plan  (GMLP Principle 10)

{monitoring_plan or "_Describe production monitoring: drift detection, performance recalibration cadence, adverse event reporting pathway, version control / Algorithm Change Protocol (ACP) plan._"}

## 9. Human Oversight  (GMLP Principle 6)

_The clinical user retains responsibility for the decision. The model's
output is advisory unless explicitly cleared as autonomous (FDA SaMD Class
III). Specify the human review touchpoint before any irreversible action._

## 10. Cybersecurity Posture Reference  (GMLP Principle 5)

A full HIPAA Security Risk Analysis derived from this same graph is available
via Cyber Nexus → Compliance → Generate SRA. CVE attack chains affecting
this model's dependencies are surfaced under Posture → Gaps & Replay.

---

_End of Model Card_
"""


def generate_docx(model_name: str = "(unspecified)", **kwargs) -> bytes:
    try:
        from docx import Document
    except ImportError:
        return (
            "# (python-docx not installed - install with `pip install python-docx`)\n\n"
            + generate_markdown(model_name, **kwargs)
        ).encode("utf-8")
    md = generate_markdown(model_name, **kwargs)
    doc = Document()
    for line in md.splitlines():
        s = line.rstrip()
        if s.startswith("# "):
            doc.add_heading(s[2:], 0)
        elif s.startswith("## "):
            doc.add_heading(s[3:], 1)
        elif s.startswith("### "):
            doc.add_heading(s[4:], 2)
        elif s.startswith("|"):
            pass  # tables are noisy in docx; keep as md text
        elif s.strip():
            doc.add_paragraph(s)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
